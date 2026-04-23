from docx import Document
from docx.shared import Pt
import os

# Create a new Document
doc = Document()

# --- Title ---
title = doc.add_heading('Serverless Object Detection Pipeline', 0)
title_run = title.runs[0]
title_run.font.bold = True

doc.add_heading('Project Overview & Real-World Applications', level=1)

# --- Executive Summary ---
p_exec = doc.add_paragraph(
    "Purpose of the Project:\n"
)
p_exec.runs[0].bold = True
p_exec.add_run(
    "This project serves as a highly scalable, serverless architecture for automated image analysis. "
    "By utilizing AWS services—most notably Amazon Rekognition, Lambda, S3, and DynamoDB—the system offloads "
    "the heavy computational burden of image processing from client devices (React Native, Android, and Web) "
    "to the cloud. The key goal is to process images optimally: safely acquiring user photos through presigned URLs, "
    "running parallel machine-learning inference (object detection, text extraction, content moderation, and "
    "checking for personal protective equipment), analyzing the data, generating annotated visual feedback, "
    "and dispatching immediate alerts (via SNS) when necessary—all without idle infrastructure overhead."
)

# --- Key Architectural Benefits ---
doc.add_heading('Key Architectural Benefits', level=2)
benefits = doc.add_paragraph()
benefits.style = 'List Bullet'
benefits.add_run("Scalability at Zero Idle Cost: ").bold = True
benefits.add_run("The serverless design means there are no servers to provision or maintain. If thousands of images process simultaneously, AWS dynamically scales. When there is zero traffic, costs drop to near zero.")

benefits2 = doc.add_paragraph()
benefits2.style = 'List Bullet'
benefits2.add_run("Low Latency Parallel Processing: ").bold = True
benefits2.add_run("Instead of doing round-trip inference searches sequentially, the backend executes labels, text (OCR), moderation, and PPE compliance concurrently, combining the final telemetry in mere milliseconds.")

benefits3 = doc.add_paragraph()
benefits3.style = 'List Bullet'
benefits3.add_run("Secure & Decoupled Uploads: ").bold = True
benefits3.add_run("Client applications use short-lived Presigned URLs, directly pushing heavy binaries straight to S3. This protects the API Gateway layer and minimizes Lambda execution limits.")

# --- Real-World Applications ---
doc.add_heading('Real-World Application Use Cases', level=1)

# Use Case 1
uc1 = doc.add_heading('1. Construction & Industrial Site Safety (Primary Use Case)', level=2)
p_uc1 = doc.add_paragraph(
    "With the built-in Personal Protective Equipment (PPE) detection module, this pipeline is natively suited for "
    "industrial safety compliance. Workers or site cameras can capture entry/exit or shift photos. The system "
    "automatically verifies if workers are wearing hard hats, safety vests, or gloves. If an individual is "
    "non-compliant, the backend triggers an SNS alert (e.g., sending an immediate email/SMS to the site manager), "
    "reducing liabilities and keeping job sites safe."
)

# Use Case 2
uc2 = doc.add_heading('2. Automated Content Moderation for Social Platforms', level=2)
p_uc2 = doc.add_paragraph(
    "For community forums, dating apps, or social media platforms where users upload massive volumes of images (like the built React Native frontend), "
    "the system automatically checks for inappropriate content (NSFW moderation). Explicit images are blocked before "
    "being publicly rendered, maintaining platform integrity and significantly reducing the workload on human moderation teams."
)

# Use Case 3
uc3 = doc.add_heading('3. Retail & Logistics Inventory Digitization', level=2)
p_uc3 = doc.add_paragraph(
    "Given the pipeline’s parallel Text Detection (OCR) and Object Labeling capabilities, warehouse operators can "
    "take photos of incoming shipments or boxes. The backend will instantly extract serial numbers, tracking codes, "
    "and recognize the physical shapes of objects (e.g., identifying damaged boxes or specific machine parts), "
    "synchronizing effortlessly with central inventory databases via DynamoDB."
)

# Use Case 4
uc4 = doc.add_heading('4. Automated Insurance Claims & Adjusting', level=2)
p_uc4 = doc.add_paragraph(
    "Policyholders submitting damage claims from their phones can upload accident photos via the Android or Web client. "
    "The Rekognition pipeline classifies vehicle parts, identifies visible damage characteristics, extracts license plates "
    "(Text OCR), and processes contextual metadata. Adjusters receive an automatically annotated (Pillow-drawn) photo "
    "and an organized dataset, massively expediting the claim verification timeline."
)

# --- Conclusion ---
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "Ultimately, this serverless object detection foundation bridges the gap between sophisticated ML vision "
    "APIs and multi-platform consumer applications. Its loosely coupled design ensures it acts as a generic "
    "funnel for any media analysis requirement—adaptable for rapid ingestion, deep analysis, real-time alerting, "
    "and durable archival logic."
)

# Save the document
filename = "Project_Purpose_And_UseCases.docx"
filepath = os.path.join(os.getcwd(), filename)
doc.save(filepath)
print(f"Document generated successfully as {filename}")